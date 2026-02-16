"""
FineTuneFlow — Text Extraction Service.

Extracts plain text from uploaded documents:
  - PDF  → PyPDF2
  - DOCX → python-docx
  - TXT/MD → direct read (UTF-8 with fallback)
"""

from __future__ import annotations

from pathlib import Path

from app.core.logging import get_logger

logger = get_logger(__name__)


def extract_text(file_path: Path, mime_type: str | None = None) -> str:
    """
    Extract text from a document file.

    Args:
        file_path: Path to the document on disk.
        mime_type: Optional MIME type hint. If None, inferred from extension.

    Returns:
        Extracted plain text content.

    Raises:
        ValueError: If the file type is unsupported or extraction fails.
        FileNotFoundError: If the file does not exist.
    """
    if not file_path.exists():
        raise FileNotFoundError(f"File not found: {file_path}")

    ext = file_path.suffix.lower()

    if ext == ".pdf" or mime_type == "application/pdf":
        return _extract_pdf(file_path)
    elif ext == ".docx" or mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        return _extract_docx(file_path)
    elif ext in (".txt", ".md") or (mime_type and mime_type.startswith("text/")):
        return _extract_text_file(file_path)
    else:
        raise ValueError(f"Unsupported file type: ext={ext}, mime={mime_type}")


def _extract_pdf(file_path: Path) -> str:
    """Extract text from a PDF file using PyPDF2."""
    from PyPDF2 import PdfReader

    logger.info("text_extractor.pdf.start", path=str(file_path))

    reader = PdfReader(str(file_path))
    pages: list[str] = []

    for i, page in enumerate(reader.pages):
        text = page.extract_text()
        if text:
            # Normalize whitespace per page
            text = text.strip()
            if text:
                pages.append(text)

    if not pages:
        logger.warning("text_extractor.pdf.empty", path=str(file_path))
        return ""

    result = "\n\n".join(pages)
    logger.info(
        "text_extractor.pdf.done",
        path=str(file_path),
        pages=len(reader.pages),
        chars=len(result),
    )
    return result


def _extract_docx(file_path: Path) -> str:
    """Extract text from a DOCX file using python-docx."""
    from docx import Document

    logger.info("text_extractor.docx.start", path=str(file_path))

    doc = Document(str(file_path))
    paragraphs: list[str] = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)

    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                paragraphs.append(" | ".join(cells))

    if not paragraphs:
        logger.warning("text_extractor.docx.empty", path=str(file_path))
        return ""

    result = "\n\n".join(paragraphs)
    logger.info(
        "text_extractor.docx.done",
        path=str(file_path),
        paragraphs=len(paragraphs),
        chars=len(result),
    )
    return result


def _extract_text_file(file_path: Path) -> str:
    """Read a plain text or markdown file with encoding fallback."""
    logger.info("text_extractor.text.start", path=str(file_path))

    # Try UTF-8 first, then fallback encodings
    # Note: latin-1 accepts any byte sequence, so it must come last
    for encoding in ("utf-8", "utf-8-sig", "cp1252", "latin-1"):
        try:
            text = file_path.read_text(encoding=encoding)
            text = text.strip()
            logger.info(
                "text_extractor.text.done",
                path=str(file_path),
                encoding=encoding,
                chars=len(text),
            )
            return text
        except (UnicodeDecodeError, UnicodeError):
            continue

    # Last resort: read as bytes and decode with errors replaced
    raw = file_path.read_bytes()
    text = raw.decode("utf-8", errors="replace").strip()
    logger.warning(
        "text_extractor.text.fallback",
        path=str(file_path),
        chars=len(text),
    )
    return text
